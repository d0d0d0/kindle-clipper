#!/usr/bin/python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Inches
from dictionary_api import *

import time

class Clipper:

	def __init__(self):
		self.clippings = None
		self.notebook = {}
		self.book_list = []
		self.dictionary = []

	def load_file(self, input):
		self.clippings = open(input, 'r', encoding='utf-8').read()
		self.create_notebook()

	def list_books(self):
		self.book_list = [book for book, _ in self.notebook.items()]

	def create_notebook(self):
		notebook = {}

		for c in filter(None, self.clippings.split("==========\n")):
			lines = list(filter(None, c.split("\n")))

			if len(lines) < 3: # there is no text
				continue 

			header, meta, text = filter(None, lines)
			
			if " " not in text: # if it is just a word
				text = re.sub(r'[^\w\s]', '', text)
				self.dictionary.append(text)
				continue

			book, writer = header[:-1].rsplit(" (", 1)
			book = re.sub('/[\x00-\x1F\x7F]/u','', book) # encode to ascii

			indicators = meta.split(" | ")

			page = 0 if len(indicators) < 3 else indicators[0].split("page ")[-1]
			date = (indicators[-1].split("Added on ")[-1]).rsplit(" ", 2)[0]

			note = {"page": page, "date": date, "text": text}

			if book not in notebook:
				notebook[book] = {"writer": writer, "notes": []}

			notebook[book]["notes"].append(note)

		self.notebook = notebook
		return notebook

	def find_notes_by_book(self, book):
		partial_notebook = {}
		for k,v in self.notebook.items():
			if book in k:
				partial_notebook[k] = v

		return partial_notebook

	def find_notes_by_writer(self, writer):
		partial_notebook = {}
		for k,v in self.notebook.items():
			if writer in v["writer"]:
				partial_notebook[k] = v

		return partial_notebook

	def build_dictionary(self):

		dictionary_text = ""

		counter = 0 # for limited api search

		start = time.time()

		self.dictionary = list(set([word.lower() for word in self.dictionary]))
		self.dictionary.sort()

		for word in self.dictionary:
			definition = search_word(word)

			if definition:
				dictionary_text += "%s: %s \n" %(word, definition)

			counter += 1

			if counter == MINUTE_LIMIT - 1 and time.time() - start < 60:
				#print("sleep started after %d count and time %f" %(counter, time.time() - start))
				time.sleep(time.time() - start + 1)
				start = time.time() 
				counter = 0

		return dictionary_text

	def write_notes(self, notebook, out, type="docx", is_dictionary=False):

		if type == "txt":
			doc = "Reading Notes\n\n"
			for k,v in notebook.items():
				doc += "%s\n%s\n\n" %(k, v["writer"])
				for n in v["notes"]:
					doc += "- %s (on p.%s at %s)\n" %(n["text"], n["page"], n["date"])
				doc += "\n"

			f = open(out + ".txt", 'w+')
			f.write(doc)
			f.close()

		else:
			document = Document()
			document.add_heading('Reading Notes', 0)
			for k,v in notebook.items():
				document.add_heading(k, level=2)
				document.add_heading(v["writer"], level=4)
				for n in v["notes"]:
					t = "%s (on p.%s at %s)\n" %(n["text"], n["page"], n["date"])
					document.add_paragraph(t, style='List Bullet')

			if is_dictionary:
				dictionary_text = self.build_dictionary()

				document.add_heading("Dictionary", level=2)
				document.add_paragraph(dictionary_text)

			document.save(out + ".docx")
